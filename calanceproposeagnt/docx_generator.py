from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
import logging
import datetime

def add_custom_styles(doc):
    # Add a custom style for TOC entries
    style = doc.styles.add_style('TOC 1', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    style.paragraph_format.space_before = Pt(6)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.left_indent = Pt(12)

def generate_docx(proposal_content, client_name):
    doc = Document()
    
    # Add custom styles
    add_custom_styles(doc)
    
    # Add default date if missing
    current_date = datetime.date.today().strftime("%Y-%m-%d")
    proposal_date = proposal_content.get("date", current_date)
    project_name = proposal_content.get("project_name", "Unnamed Project")

    # Log warnings for missing key information
    if "date" not in proposal_content:
        logging.warning("Date is missing from proposal content")
    if "project_name" not in proposal_content:
        logging.warning("Project name is missing from proposal content")

    # Add a title with error handling
    try:
        doc.add_heading(f'SOW {proposal_date} {project_name} for {client_name}', 0)
    except Exception as e:
        logging.error(f"Error adding heading: {str(e)}")
        doc.add_heading(f'SOW for {client_name}', 0)

    # Add Calance information
    doc.add_paragraph('Calance\n888 Disneyland Drive Suite 500\nAnaheim, CA 92802')

    # Add a table for date, services performed by, and services performed for
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Date'
    hdr_cells[1].text = 'Services Performed By:'
    hdr_cells[2].text = 'Services Performed For:'
    row_cells = table.rows[1].cells
    row_cells[0].text = proposal_date
    row_cells[1].text = 'Calance\n888 Disneyland Drive Suite 500\nAnaheim, CA 92802'
    row_cells[2].text = f"{client_name}\n[Client Address]\n[City, ST ZIP Code]"

    # Add table of contents
    doc.add_heading('Table of Contents', level=1)
    for section in proposal_content.keys():
        if section not in ['date', 'project_name']:
            doc.add_paragraph(section.replace('_', ' ').title(), style='TOC 1')

    # Add sections from the proposal content with error handling
    for section in ['description', 'purpose', 'scope', 'approach', 'engagement_approach', 'project_estimated_timeline', 'development_hosting_support_maintenance_estimates', 'risks_constraints_dependencies']:
        try:
            doc.add_heading(section.replace('_', ' ').title(), level=1)
            content = proposal_content.get(section, "Not provided")
            doc.add_paragraph(str(content))
        except Exception as e:
            logging.error(f"Error adding section {section}: {str(e)}")

    # Add Terms and Conditions
    doc.add_heading('Terms and Conditions', level=1)
    terms = [
        'The pricing in this SOW is valid for 30 days from the date of submission.',
        'Professional fees will be invoiced monthly based on actual time & materials effort.',
        'We do not anticipate travel for this project.'
    ]
    for term in terms:
        doc.add_paragraph(term, style='ListBullet')

    # Add Signatures / Work Authorization
    doc.add_heading('Signatures / Work Authorization', level=1)
    doc.add_paragraph('IN WITNESS WHEREOF, and in acknowledgment that the parties hereto have read and understood each and every provision hereof, the parties have executed this Statement of Work.')

    # Add signature lines
    doc.add_paragraph(f"Client: {client_name}")
    doc.add_paragraph("By: ________________________________")
    doc.add_paragraph("Name: ______________________________")
    doc.add_paragraph("Title: _______________________________")
    doc.add_paragraph("Date: _______________________________")

    doc.add_paragraph("Vendor: Partners Information Technology, Inc. dba Calance")
    doc.add_paragraph("By: ________________________________")
    doc.add_paragraph("Name: Asit Govil")
    doc.add_paragraph("Title: COO")
    doc.add_paragraph("Date: _______________________________")

    # Add headers and footers
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = f"SOW for {client_name}"
    footer = section.footer
    footer.paragraphs[0].text = "Page "

    return doc
